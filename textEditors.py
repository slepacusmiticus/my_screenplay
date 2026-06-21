import os
from config import APP_NAME
from PyQt6.QtCore import Qt, QSettings
from PyQt6.QtGui import QTextDocument, QTextDocumentWriter
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QLabel,
    QTabWidget, QTabBar, QSlider, QInputDialog,
    QFileDialog, QMessageBox
)


# ── Document registry ─────────────────────────────────────────────────────────
# Maps (editor_class, label) → QTextDocument so that the same label opened in
# two different panels of the same editor type shares one live document.
# Edits in either panel are reflected immediately — no polling or prompts.
# Keys are (type, str) so NoteEditor "Untitled 1" ≠ ScreenplayEditor "Untitled 1".
_DOC_REGISTRY: dict[tuple, QTextDocument] = {}


# ── TabbedEditor ──────────────────────────────────────────────────────────────
# Shared base class for all tabbed text editors.
# Provides: tabbed pages, permanent first tab, incremented Untitled N naming,
# scale slider status bar, and file save/open/close operations.
# Files are saved as HTML (.note) so rich formatting is fully preserved.
# Subclasses only need to set `name` for the Panel dropdown label.

class TabbedEditor(QWidget):
    name = ""
    _BASE_FONT_SIZE = 12   # point size at 100 %
    _FILE_FILTER    = "Note files (*.note);;All files (*)"

    def __init__(self):
        super().__init__()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self._tabs = QTabWidget()
        self._tabs.setTabsClosable(True)
        self._tabs.tabCloseRequested.connect(self._close_tab)
        self._tabs.tabBarDoubleClicked.connect(self._on_tab_double_clicked)
        layout.addWidget(self._tabs)

        self._untitled_counter = 0   # increments with each new tab, never resets
        self._paths: dict[QTextEdit, str] = {}  # editor widget → saved file path

        # ── Status bar ───────────────────────────────────────────────────────
        status = QWidget()
        status.setFixedHeight(24)
        status_layout = QHBoxLayout(status)
        status_layout.setContentsMargins(6, 0, 6, 0)
        status_layout.setSpacing(6)

        status_layout.addStretch()
        status_layout.addWidget(QLabel("Scale:"))

        self._scale_slider = QSlider(Qt.Orientation.Horizontal)
        self._scale_slider.setRange(50, 200)
        self._scale_slider.setValue(100)
        self._scale_slider.setFixedWidth(120)
        self._scale_slider.setTickInterval(50)
        self._scale_slider.valueChanged.connect(self._on_scale_changed)
        status_layout.addWidget(self._scale_slider)

        self._scale_label = QLabel("100%")
        self._scale_label.setFixedWidth(36)
        status_layout.addWidget(self._scale_label)

        layout.addWidget(status)

        # Create the initial permanent tab and strip its close button.
        self._add_tab(closable=False)

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _add_tab(self, closable=True):
        self._untitled_counter += 1
        label = f"Untitled {self._untitled_counter}"

        # Share one QTextDocument per (editor type, label) so that the same
        # document open in multiple panels stays in sync automatically.
        key = (type(self), label)
        if key not in _DOC_REGISTRY:
            _DOC_REGISTRY[key] = QTextDocument()
        doc = _DOC_REGISTRY[key]

        editor = QTextEdit()
        editor.setDocument(doc)
        font = editor.font()
        font.setPointSize(self._current_font_size())
        editor.setFont(font)
        idx = self._tabs.addTab(editor, label)
        if not closable:
            bar = self._tabs.tabBar()
            assert bar is not None
            bar.setTabButton(idx, QTabBar.ButtonPosition.RightSide, None)
        self._tabs.setCurrentIndex(idx)

    def _close_tab(self, index):
        widget = self._tabs.widget(index)
        if isinstance(widget, QTextEdit):
            self._paths.pop(widget, None)
        self._tabs.removeTab(index)

    def _on_tab_double_clicked(self, index: int) -> None:
        current = self._tabs.tabText(index)
        name, ok = QInputDialog.getText(self, "Rename Tab", "Name:", text=current)
        if ok and name.strip():
            self._tabs.setTabText(index, name.strip())

    def _current_editor(self) -> QTextEdit | None:
        w = self._tabs.currentWidget()
        return w if isinstance(w, QTextEdit) else None

    def _current_font_size(self):
        return max(6, int(self._BASE_FONT_SIZE * self._scale_slider.value() / 100))

    def _on_scale_changed(self, value):
        self._scale_label.setText(f"{value}%")
        size = max(6, int(self._BASE_FONT_SIZE * value / 100))
        for i in range(self._tabs.count()):
            widget = self._tabs.widget(i)
            if isinstance(widget, QTextEdit):
                font = widget.font()
                font.setPointSize(size)
                widget.setFont(font)

    def _default_dir(self) -> str:
        """Return the workspace folder set in UI Settings, or empty string."""
        return QSettings(APP_NAME, APP_NAME).value("workspace_folder", "")

    def _write(self, editor: QTextEdit, path: str) -> bool:
        """Write editor content as HTML to path. Returns True on success."""
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(editor.toHtml())
            return True
        except OSError as e:
            QMessageBox.warning(self, "Save Failed", str(e))
            return False

    # ── File operations ───────────────────────────────────────────────────────

    def new_file(self):
        self._add_tab(closable=True)

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Open File", self._default_dir(), self._FILE_FILTER
        )
        if not path:
            return
        try:
            with open(path, encoding="utf-8") as f:
                html = f.read()
        except OSError as e:
            QMessageBox.warning(self, "Open Failed", str(e))
            return

        self._add_tab(closable=True)
        editor = self._current_editor()
        if editor:
            editor.setHtml(html)
            self._paths[editor] = path
            self._tabs.setTabText(
                self._tabs.currentIndex(),
                os.path.splitext(os.path.basename(path))[0]
            )

    def save_file(self):
        editor = self._current_editor()
        if editor and editor in self._paths:
            self._write(editor, self._paths[editor])
        else:
            self.save_file_as()

    def save_file_as(self):
        editor = self._current_editor()
        if not editor:
            return
        default_name = self._tabs.tabText(self._tabs.currentIndex())
        if not default_name.endswith(".note"):
            default_name += ".note"
        default_path = os.path.join(self._default_dir(), default_name)
        path, _ = QFileDialog.getSaveFileName(
            self, "Save File As", default_path, self._FILE_FILTER
        )
        if not path:
            return
        if not path.endswith(".note"):
            path += ".note"
        if self._write(editor, path):
            self._paths[editor] = path
            self._tabs.setTabText(
                self._tabs.currentIndex(),
                os.path.splitext(os.path.basename(path))[0]
            )

    def export_as(self, fmt: str) -> None:
        """Export current tab to fmt: 'txt', 'odt', 'rtf', or 'docx'."""
        editor = self._current_editor()
        if not editor:
            return

        tab_name = self._tabs.tabText(self._tabs.currentIndex())

        filters = {
            "txt":  ("Export as Plain Text",       f"{tab_name}.txt",  "Text files (*.txt);;All files (*)"),
            "odt":  ("Export as Open Document",    f"{tab_name}.odt",  "Open Document (*.odt);;All files (*)"),
            "rtf":  ("Export as Rich Text",        f"{tab_name}.rtf",  "Rich Text (*.rtf);;All files (*)"),
            "docx": ("Export as Word Document",    f"{tab_name}.docx", "Word Document (*.docx);;All files (*)"),
        }
        title, default_name, file_filter = filters[fmt]
        default_path = os.path.join(self._default_dir(), default_name)

        path, _ = QFileDialog.getSaveFileName(self, title, default_path, file_filter)
        if not path:
            return
        if not path.endswith(f".{fmt}"):
            path += f".{fmt}"

        if fmt == "txt":
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(editor.toPlainText())
            except OSError as e:
                QMessageBox.warning(self, "Export Failed", str(e))

        elif fmt == "odt":
            writer = QTextDocumentWriter(path, b"ODF")
            if not writer.write(editor.document()):
                QMessageBox.warning(self, "Export Failed", "Could not write ODF file.")

        elif fmt == "rtf":
            # Qt6 dropped native RTF support; generate a minimal valid RTF file.
            text = editor.toPlainText()
            safe = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
            body = "\\par\n".join(safe.split("\n"))
            rtf = "{\\rtf1\\ansi\\deff0 " + body + "}"
            try:
                with open(path, "w", encoding="ascii", errors="replace") as f:
                    f.write(rtf)
            except OSError as e:
                QMessageBox.warning(self, "Export Failed", str(e))

        elif fmt == "docx":
            try:
                import docx  # python-docx
                doc = docx.Document()
                for line in editor.toPlainText().split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
            except ImportError:
                QMessageBox.warning(
                    self, "Export Failed",
                    "python-docx is not installed.\n\nInstall it with:\n  pip install python-docx"
                )
            except Exception as e:
                QMessageBox.warning(self, "Export Failed", str(e))

    def close_current_tab(self):
        idx = self._tabs.currentIndex()
        if idx > 0:   # tab 0 is permanent
            self._close_tab(idx)


# ── Concrete tabbed editors ───────────────────────────────────────────────────
# Full TabbedEditor subclasses — add editor-specific behaviour here as needed.

class NoteEditor(TabbedEditor):
    name = "Note"


class NovelEditor(TabbedEditor):
    name = "Novel"


class ScratchPadEditor(TabbedEditor):
    name = "Scratch Pad"


# ── Simple stub editors ───────────────────────────────────────────────────────
# Plain QTextEdit wrappers — placeholders until each gets its own implementation.

class TimelineEditor(QTextEdit):
    name = "Timeline Editor"


class ScratchBoardEditor(QTextEdit):
    name = "Scratch Board"


class CharacterEditor(QTextEdit):
    name = "Character Editor"
